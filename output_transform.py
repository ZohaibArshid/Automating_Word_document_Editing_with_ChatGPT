
import re
from docx import Document



def transform(section1,section2,section5):
    section1 = section1.strip()
    section2 = section2.strip()
    section5 = section5.strip()
    # fetching required numbering from gpt response
    section1_cleaned =  "\n".join(re.findall(r'1\.1\.[1-4]\..*', section1))
    section2_cleaned = "\n".join(re.findall(r'2\.1\.[1-4]\..*', section2))
    section5_cleaned = "\n".join(re.findall(r'5\.1\.(?:[1-9]|1[0-6])\..*', section5))
    # converting cleaned text into a usable list
    line_1 = section1_cleaned.split('\n')
    line_2 = section2_cleaned.split('\n')
    line_3 = section5_cleaned.split('\n')
    # # Remove the numbering from each line
    line_1 = [re.sub(r'^\d+\.\d+\.\d+\.\s*', '', line) for line in line_1]
    line_2 = [re.sub(r'^\d+\.\d+\.\d+\.\s*', '', line) for line in line_2]
    line_3 = [re.sub(r'^\d+\.\d+\.\d+\.\s*', '', line) for line in line_3]
    #removing empty spaces, incase if any #check
    line_1 = [item for item in line_1 if item != '']
    line_2 = [item for item in line_2 if item != '']
    line_3 = [item for item in line_3 if item != '']

    return line_1,line_2,line_3


def edit_doc(word_file_path,save_path,line_1,line_2,line_3):
    doc = Document(word_file_path)
    changes = {
                30: line_1[0],
                32: line_1[1],
                34: line_1[2],
                36: line_1[3],
                40: line_2[0],
                42: line_2[1],
                44: line_2[2],
                46: line_2[3],
                171:line_3[0],
                173:line_3[1],
                175:line_3[2],
                179:line_3[3],
                181:line_3[4],
                183:line_3[5],
                185:line_3[6],
                189:line_3[7],
                191:line_3[8],
                195:line_3[9],
                197:line_3[10],
                199:line_3[11],
                201:line_3[12],
                205:line_3[13],
                207:line_3[14],
                209:line_3[15],

              }
     # Loop over the dictionary
    for index, new_text in changes.items():
        # Clear all runs from the paragraph
        for run in doc.paragraphs[index].runs:
            run.text = ''
        doc.paragraphs[index].add_run(new_text)
    
        # Save the final document with all changes
    doc.save(save_path)





